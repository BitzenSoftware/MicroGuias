# -*- coding: utf-8 -*-
"""Gera 1 .docx por tema, cada um com 10 prompts de ebook + prompts de capa."""
import os
from docx import Document
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Templates (mesmos da coleção)
# ---------------------------------------------------------------------------
def ebook_prompt(e):
    caps = "\n".join(f"{i+3}. Capítulo {i+1} — {c}" for i, c in enumerate(e["capitulos"]))
    return f'''Crie um ebook prático e direto chamado "{e['titulo']}: {e['subtitulo']}".

Público-alvo: {e['publico']}

Tom: motivador, simples e direto. Linguagem fácil (português do Brasil), frases curtas, zero jargão. Trate o leitor por "você".

Formato: ebook de 12 a 15 páginas, objetivo e acionável. Cada capítulo cabe em 1 página, com no máximo 3 parágrafos curtos, e termina com um quadro destacado "{e['destaque']}" — uma tarefa ou dica simples que o leitor pode aplicar na hora.

Estrutura:
1. Capa atraente com o título e subtítulo
2. Introdução: {e['intro']}
{caps}
10. Conclusão + {e['conclusao']}

Design: {e['design']}. Inclua um quadro de destaque "{e['destaque']}" em cada capítulo.

Adicione no rodapé: "Micro Guias".'''


def capa_prompt(e):
    t = e["teasers"]
    return f'''Create ONE single professional 3D ebook cover mockup. Output must contain ONLY ONE book — never multiple books, never variations, never a collage.

BOOK / FORMAT:
- A single hardcover book of normal thickness, standing upright in 3/4 perspective, rotated about 20 degrees to the RIGHT (spine and page edges visible on the RIGHT side)
- The ENTIRE book and the FULL title fully visible and centered, with comfortable empty margin — do NOT crop anything
- Clean solid white background, soft realistic shadow and a subtle reflection below
- Photorealistic, high quality, premium luxury collection style

COVER DESIGN (top to bottom):
1. A top pill badge with a thin golden outline containing the text "EDIÇÃO ESPECIAL" in light letters
2. A clean, organized 2x2 GRID of four separate rectangular photos, each with a thin golden border, aligned in perfect rows and columns — NOT floating, NOT overlapping, NOT scattered. The four photos: {e['fotos']}
3. A central rounded banner in DARK deep navy blue color, framed by a thin elegant GOLDEN ornamental border; inside it the main title in a refined golden serif font
4. Just below, the subtitle in a clean light/golden font
5. A dark bottom band with 4 short teasers in small golden uppercase letters, separated by dots

TEXTS (exactly, in Portuguese):
- Title: "{e['titulo']}"
- Subtitle: "{e['subtitulo']}"
- Teasers: "{t[0]}" • "{t[1]}" • "{t[2]}" • "{t[3]}"

COLORS: {e['cores_capa']}; the title banner and bottom band are always dark navy with golden details. Theme accent colors only inside the photos.

OUTPUT: a single vertical book cover, 2:3 aspect ratio, one book only, centered, full book visible, with a CLEAN ALIGNED 2x2 photo grid, premium and elegant.'''


def E(titulo, subtitulo, publico, intro, capitulos, conclusao, design, fotos, teasers, cores_capa, destaque="Ação de hoje"):
    return dict(titulo=titulo, subtitulo=subtitulo, publico=publico, intro=intro,
                capitulos=capitulos, conclusao=conclusao, design=design, fotos=fotos,
                teasers=teasers, cores_capa=cores_capa, destaque=destaque)

# ===========================================================================
# FINANÇAS
# ===========================================================================
FINANCAS = [
    E("Saia das Dívidas em 90 Dias", "O Plano Para Limpar Seu Nome Sem Sufoco",
      "pessoas endividadas e no negativo que querem um plano realista para zerar as dívidas e voltar a respirar.",
      "por que você não consegue sair das dívidas (e a verdade que ninguém conta)",
      ["Coloque todas as dívidas na mesa", "A ordem certa para pagar (bola de neve vs avalanche)",
       "Como negociar e conseguir descontos", "Corte gastos invisíveis sem sofrer",
       "Renegocie juros e parcelas", "Crie uma renda extra para acelerar", "Os erros que mantêm você endividado"],
      "checklist 'meu plano de 90 dias para sair do vermelho'",
      "visual sério e acolhedor, tons de azul e verde, ícones de finanças, fácil de ler no celular",
      "a person reviewing bills with relief, a calculator and notebook, a happy person free of debt, a financial planning desk",
      ["A ORDEM CERTA DE PAGAR", "NEGOCIE COM DESCONTO", "CORTE GASTOS INVISÍVEIS", "PLANO DE 90 DIAS"],
      "tones of blue and green"),
    E("Comece a Investir do Zero", "O Guia Para Quem Nunca Investiu na Vida",
      "iniciantes totais que querem investir mas têm medo, acham complicado ou pensam que precisam de muito dinheiro.",
      "por que deixar dinheiro parado é perder dinheiro todo mês",
      ["Antes de investir: o que organizar primeiro", "Onde abrir sua conta para investir",
       "Renda fixa: o melhor lugar para começar", "Tesouro Direto explicado de forma simples",
       "O que são ações e fundos (sem complicar)", "Quanto investir todo mês (mesmo com pouco)", "Os erros de iniciante que custam caro"],
      "checklist 'seu primeiro investimento esta semana'",
      "visual moderno e confiável, tons de verde e azul, ícones de crescimento, fácil de ler no celular",
      "a beginner investing on a phone app, growth charts going up, coins and a piggy bank, a confident young investor",
      ["ORGANIZE ANTES", "RENDA FIXA SIMPLES", "TESOURO DIRETO", "COMECE COM POUCO"],
      "tones of green and blue"),
    E("Reserva de Emergência", "Como Montar Seu Colchão de Segurança do Zero",
      "pessoas que vivem no aperto e querem criar uma reserva para imprevistos sem passar sufoco.",
      "por que quem não tem reserva vive refém de qualquer imprevisto",
      ["Quanto você realmente precisa guardar", "Onde deixar a reserva (com liquidez)",
       "Como começar guardando pouco", "Automatize e não dependa de força de vontade",
       "Como cortar gastos para acelerar", "O que NÃO é emergência", "Os erros que esvaziam sua reserva"],
      "checklist 'sua reserva começa hoje'",
      "visual seguro e clean, tons de azul e verde-água, ícones de proteção, fácil de ler no celular",
      "a safe with savings, a person feeling secure, a piggy bank and shield icon, an organized budget on a phone",
      ["QUANTO GUARDAR", "ONDE DEIXAR", "AUTOMATIZE", "COMECE COM POUCO"],
      "tones of blue and teal"),
    E("Organize suas Finanças em 30 Dias", "Do Caos ao Controle Total do Seu Dinheiro",
      "pessoas que não sabem para onde vai o dinheiro e querem assumir o controle das finanças.",
      "por que você termina o mês sem saber onde gastou",
      ["Descubra para onde seu dinheiro vai", "Monte um orçamento que funciona de verdade",
       "A regra 50-30-20 simplificada", "Separe contas, metas e lazer", "Apps e planilhas para controlar tudo",
       "Crie metas financeiras realistas", "Os erros que sabotam seu orçamento"],
      "checklist 'finanças organizadas em 30 dias'",
      "visual organizado e moderno, tons de azul e laranja, ícones de planejamento, fácil de ler no celular",
      "an organized budget spreadsheet, a person planning finances, a calendar with money goals, a tidy financial desk",
      ["PARA ONDE VAI SEU DINHEIRO", "ORÇAMENTO QUE FUNCIONA", "REGRA 50-30-20", "METAS REALISTAS"],
      "tones of blue and orange"),
    E("Renda Extra", "15 Formas Reais de Ganhar Mais Sem Sair de Casa",
      "pessoas que querem complementar a renda com formas práticas e acessíveis, mesmo com pouco tempo.",
      "por que depender de um único salário é arriscado",
      ["Venda o que você já sabe fazer", "Serviços digitais que dão dinheiro",
       "Revenda e dropshipping para iniciantes", "Produtos digitais (ebooks, cursos)",
       "Trabalhos freelancer mais procurados", "Como divulgar e conseguir os primeiros clientes", "Os erros que fazem desistir cedo"],
      "lista 'escolha sua renda extra e comece em 7 dias'",
      "visual aspiracional e prático, tons de verde, dourado e azul, ícones de dinheiro, fácil de ler no celular",
      "a person earning extra income at home, a laptop with online work, money growth, a freelancer working happily",
      ["SERVIÇOS DIGITAIS", "PRODUTOS DIGITAIS", "FREELANCER", "COMECE EM 7 DIAS"],
      "tones of green, gold and blue"),
    E("Como Sair do Vermelho e Poupar", "O Passo a Passo Para Virar o Jogo",
      "pessoas no negativo todo mês que querem equilibrar as contas e começar a sobrar dinheiro.",
      "a diferença entre quem vive no vermelho e quem prospera",
      ["Pare o sangramento: gastos que te afundam", "Aumente o que entra, reduza o que sai",
       "Renegocie cartão e cheque especial", "O método para sobrar no fim do mês",
       "Construa o hábito de poupar", "Metas de curto e longo prazo", "Os erros que mantêm você no vermelho"],
      "checklist 'do vermelho ao azul em 60 dias'",
      "visual motivador, tons de azul e verde, ícones de virada e crescimento, fácil de ler no celular",
      "a person turning finances around, red to green chart, a relieved face, savings growing",
      ["PARE O SANGRAMENTO", "RENEGOCIE DÍVIDAS", "SOBRE NO FIM DO MÊS", "VIRE O JOGO"],
      "tones of blue and green"),
    E("Planejamento Financeiro Familiar", "Como Cuidar do Dinheiro da Casa em Equipe",
      "famílias e casais que querem organizar as contas da casa, evitar brigas por dinheiro e realizar sonhos juntos.",
      "por que dinheiro é uma das maiores causas de briga em casa",
      ["Junte as contas da casa numa visão só", "Orçamento familiar que todos respeitam",
       "Como dividir despesas de forma justa", "Metas em família (viagem, casa, escola)",
       "Ensine as crianças sobre dinheiro", "Reserva e proteção para a família", "Os erros que desorganizam as contas da casa"],
      "checklist 'finanças da família organizadas'",
      "visual acolhedor e familiar, tons de azul e verde, ícones de casa e família, fácil de ler no celular",
      "a family planning finances together, a couple reviewing a budget, a home and savings, a happy family",
      ["ORÇAMENTO FAMILIAR", "DIVISÃO JUSTA", "METAS EM FAMÍLIA", "PROTEÇÃO DA CASA"],
      "tones of blue and green"),
    E("Independência Financeira", "O Primeiro Passo Para Viver de Renda",
      "pessoas que sonham em ter liberdade financeira e querem entender por onde começar de forma realista.",
      "o que realmente significa ser livre financeiramente",
      ["Quanto você precisa para ser livre", "A fórmula simples da independência",
       "Aumente a distância entre ganhos e gastos", "Faça seu dinheiro trabalhar por você",
       "Os investimentos que geram renda", "Acelere com renda extra e foco", "Os erros que adiam sua liberdade"],
      "checklist 'seu mapa para a liberdade financeira'",
      "visual inspirador e sofisticado, tons de azul-marinho e dourado, ícones de liberdade, fácil de ler no celular",
      "a person enjoying financial freedom, passive income concept, money working, a relaxed lifestyle",
      ["QUANTO VOCÊ PRECISA", "A FÓRMULA SIMPLES", "RENDA PASSIVA", "SEU MAPA DA LIBERDADE"],
      "tones of navy blue and gold"),
    E("Como Economizar no Dia a Dia", "100 Formas de Gastar Menos Sem Abrir Mão de Viver",
      "pessoas que querem economizar nas pequenas coisas e ver a diferença sobrar no fim do mês.",
      "como pequenas economias viram grandes resultados",
      ["Economize nas compras do mercado", "Corte gastos com contas de casa (luz, água, internet)",
       "Transporte e combustível: gaste menos", "Lazer barato sem deixar de se divertir",
       "Compras conscientes: pare de gastar por impulso", "Aplicativos que ajudam a economizar", "Os erros que furam seu orçamento"],
      "checklist 'economias que cabem na sua rotina'",
      "visual prático e leve, tons de verde e azul, ícones de economia, fácil de ler no celular",
      "a person saving money on groceries, a shopping cart with savings, coins in a jar, smart spending",
      ["ECONOMIA NO MERCADO", "CONTAS DE CASA", "COMPRE CONSCIENTE", "100 FORMAS DE POUPAR"],
      "tones of green and blue"),
    E("Educação Financeira para Iniciantes", "Tudo Que a Escola Não Te Ensinou Sobre Dinheiro",
      "pessoas que nunca aprenderam a lidar com dinheiro e querem dominar o básico de uma vez por todas.",
      "por que ninguém te ensinou a cuidar do seu dinheiro",
      ["A mentalidade que muda sua vida financeira", "Como o dinheiro realmente funciona",
       "Gastar, poupar e investir: o equilíbrio", "Fuja das armadilhas do consumo e crédito",
       "Juros: seu amigo ou seu inimigo", "Construa patrimônio do zero", "Os erros financeiros mais comuns"],
      "checklist 'seus primeiros passos com o dinheiro'",
      "visual educativo e moderno, tons de azul e verde, ícones de aprendizado, fácil de ler no celular",
      "a person learning about money, financial education concept, a book and coins, a confident learner",
      ["MENTALIDADE FINANCEIRA", "GASTAR x POUPAR x INVESTIR", "ARMADILHAS DO CRÉDITO", "PRIMEIROS PASSOS"],
      "tones of blue and green"),
]

# ===========================================================================
# SAÚDE & BEM-ESTAR
# ===========================================================================
SAUDE = [
    E("Emagreça Sem Dietas Malucas", "O Método Realista Para Perder Peso e Manter",
      "pessoas cansadas de dietas restritivas que querem emagrecer de forma sustentável e sem sofrimento.",
      "por que dietas radicais sempre falham (e o que funciona de verdade)",
      ["O segredo do déficit calórico sem passar fome", "Monte um prato que sacia e emagrece",
       "Os alimentos que trabalham a seu favor", "Como vencer a fome emocional",
       "Movimento simples que acelera resultados", "Crie hábitos que se mantêm", "Os erros que travam o emagrecimento"],
      "checklist 'seu plano realista de emagrecimento'",
      "visual leve e motivador, tons de verde e azul, ícones de saúde, fácil de ler no celular",
      "a healthy balanced plate, a person feeling lighter, fresh vegetables, a happy fit lifestyle",
      ["SEM PASSAR FOME", "PRATO QUE SACIA", "FOME EMOCIONAL", "PLANO REALISTA"],
      "tones of green and blue"),
    E("7 Dias Para Desinchar", "Adeus Inchaço e Retenção de Líquido",
      "pessoas que se sentem inchadas e pesadas e querem desinchar rápido de forma natural.",
      "por que você acorda inchado e o que muda isso",
      ["Os vilões que causam inchaço", "Alimentos que desincham naturalmente",
       "A importância da água (e como beber mais)", "Reduza o sódio escondido",
       "Movimente o corpo para drenar", "Sono e estresse: o elo com o inchaço", "Os erros que pioram a retenção"],
      "plano '7 dias para desinchar' passo a passo",
      "visual fresco e clean, tons de verde-água e branco, ícones de leveza, fácil de ler no celular",
      "a person feeling light and healthy, fresh water and vegetables, a flat belly concept, a refreshing healthy scene",
      ["VILÕES DO INCHAÇO", "ALIMENTOS QUE DESINCHAM", "MENOS SÓDIO", "PLANO DE 7 DIAS"],
      "tones of teal and white"),
    E("Comece a Correr", "Do Sofá aos 5km em 6 Semanas",
      "pessoas sedentárias que querem começar a correr do zero, sem se machucar e sem desistir.",
      "por que correr é o exercício mais democrático que existe",
      ["Antes de correr: prepare o corpo", "O método caminhar-correr para iniciantes",
       "Tênis e roupa: o básico que importa", "Respiração e ritmo certos",
       "Evite lesões e dores comuns", "O plano de 6 semanas até os 5km", "Os erros que fazem iniciantes parar"],
      "plano 'do sofá aos 5km' semana a semana",
      "visual energético e jovem, tons de azul e laranja, ícones de corrida, fácil de ler no celular",
      "a person running outdoors, running shoes, a beginner jogger smiling, a sunny running track",
      ["CAMINHAR E CORRER", "EVITE LESÕES", "RESPIRAÇÃO CERTA", "PLANO DE 6 SEMANAS"],
      "tones of blue and orange"),
    E("Alimentação Saudável para Iniciantes", "Como Comer Melhor Sem Complicar",
      "pessoas que querem comer de forma mais saudável mas se perdem em tanta informação.",
      "por que comer bem é mais simples do que te vendem",
      ["Monte o prato saudável perfeito", "Substituições inteligentes que você nem percebe",
       "Lanches saudáveis que matam a vontade", "Como ler rótulos sem cair em pegadinhas",
       "Planeje as refeições da semana", "Coma melhor gastando menos", "Os erros que sabotam sua alimentação"],
      "checklist 'sua semana de alimentação saudável'",
      "visual fresco e apetitoso, tons de verde e amarelo, ícones de comida saudável, fácil de ler no celular",
      "a colorful healthy meal, fresh fruits and vegetables, a person preparing healthy food, a balanced plate",
      ["O PRATO PERFEITO", "SUBSTITUIÇÕES ESPERTAS", "LEIA OS RÓTULOS", "SEMANA SAUDÁVEL"],
      "tones of green and yellow"),
    E("Controle a Ansiedade Naturalmente", "Técnicas Simples Para Acalmar a Mente",
      "pessoas que sofrem com ansiedade no dia a dia e querem ferramentas práticas para se acalmar.",
      "o que é a ansiedade e por que ela aparece",
      ["A respiração que desliga o alerta", "Como parar os pensamentos acelerados",
       "O poder do agora: atenção plena simples", "Movimento e natureza contra a ansiedade",
       "Reduza estímulos: telas, café e notícias", "Crie uma rotina que acalma", "Quando procurar ajuda profissional"],
      "checklist 'sua caixa de ferramentas anti-ansiedade'",
      "visual calmo e acolhedor, tons de azul e lilás, ícones de bem-estar, fácil de ler no celular",
      "a calm person breathing, a peaceful nature scene, meditation concept, a relaxed mind",
      ["RESPIRAÇÃO QUE ACALMA", "PARE OS PENSAMENTOS", "ATENÇÃO PLENA", "ROTINA QUE ACALMA"],
      "tones of blue and lavender"),
    E("Mais Energia", "Vença o Cansaço e Tenha Disposição o Dia Todo",
      "pessoas que vivem cansadas e sem energia e querem recuperar a disposição de forma natural.",
      "por que você está sempre cansado mesmo dormindo",
      ["Os ladrões de energia do seu dia", "Alimente-se para ter mais disposição",
       "Hidratação e energia: a ligação esquecida", "Movimento que recarrega (e não cansa)",
       "Sono de qualidade para acordar bem", "Pausas estratégicas que renovam", "Os erros que drenam sua energia"],
      "checklist 'seu dia com mais energia'",
      "visual vibrante e positivo, tons de laranja e verde, ícones de energia, fácil de ler no celular",
      "an energetic person in the morning, healthy energizing food, sunrise and vitality, an active happy lifestyle",
      ["LADRÕES DE ENERGIA", "COMA PARA TER DISPOSIÇÃO", "MOVIMENTO QUE RECARREGA", "MAIS ENERGIA HOJE"],
      "tones of orange and green"),
    E("Adeus Dor nas Costas", "Postura e Alívio Para a Vida Moderna",
      "pessoas que sentem dores nas costas e no pescoço por má postura, trabalho sentado ou celular.",
      "por que sua coluna sofre na vida moderna",
      ["Descubra a causa da sua dor", "A postura correta sentado e em pé",
       "Ajuste sua mesa e seu celular", "Alongamentos que aliviam na hora",
       "Fortaleça o core para proteger a coluna", "Hábitos que poupam suas costas", "Os erros que pioram a dor"],
      "checklist 'sua rotina anti-dor nas costas'",
      "visual clean e saudável, tons de azul e verde, ícones de corpo e postura, fácil de ler no celular",
      "a person with good posture, a stretching exercise, an ergonomic desk setup, a healthy spine concept",
      ["A CAUSA DA DOR", "POSTURA CORRETA", "ALONGAMENTOS QUE ALIVIAM", "ROTINA ANTI-DOR"],
      "tones of blue and green"),
    E("Hábitos Saudáveis em 21 Dias", "Pequenas Mudanças, Grandes Resultados",
      "pessoas que querem adotar hábitos saudáveis de verdade e parar de desistir na primeira semana.",
      "por que a força de vontade não basta para criar hábitos",
      ["Como um hábito realmente se forma", "Comece ridiculamente pequeno",
       "Encadeie hábitos no que você já faz", "Crie gatilhos e elimine fricções",
       "Não quebre a corrente: o poder da constância", "Recupere-se quando falhar", "Os erros que matam seus hábitos"],
      "plano 'monte seus 3 hábitos em 21 dias'",
      "visual motivador e organizado, tons de verde e azul, ícones de calendário e check, fácil de ler no celular",
      "a habit tracker calendar, a person building routines, healthy daily habits, a motivated lifestyle",
      ["COMECE PEQUENO", "ENCADEIE HÁBITOS", "NÃO QUEBRE A CORRENTE", "PLANO DE 21 DIAS"],
      "tones of green and blue"),
    E("Jejum Intermitente para Iniciantes", "O Guia Simples Para Começar com Segurança",
      "pessoas curiosas sobre jejum intermitente que querem começar do jeito certo, sem riscos.",
      "o que é o jejum intermitente (sem mitos)",
      ["Como o jejum funciona no seu corpo", "Os principais protocolos (16:8, 14:10)",
       "O que pode e o que quebra o jejum", "Como começar sem passar mal",
       "O que comer na janela de alimentação", "Jejum e exercício: como combinar", "Quem não deve fazer jejum"],
      "checklist 'seu primeiro jejum com segurança'",
      "visual clean e moderno, tons de azul e verde, ícones de relógio e alimentação, fácil de ler no celular",
      "a clock and healthy food concept, a person with a balanced meal, intermittent fasting schedule, a healthy lifestyle",
      ["COMO FUNCIONA", "PROTOCOLO 16:8", "O QUE QUEBRA O JEJUM", "COMECE COM SEGURANÇA"],
      "tones of blue and green"),
    E("Digestão Leve", "Adeus Azia, Gases e Má Digestão",
      "pessoas que sofrem com azia, gases e desconforto após comer e querem uma digestão tranquila.",
      "por que sua digestão anda pesada",
      ["Os alimentos que pesam na digestão", "Mastigue e coma do jeito certo",
       "Os campeões da boa digestão", "Combata azia e refluxo naturalmente",
       "Reduza gases e desconforto", "Rotina e intestino funcionando", "Os erros que estragam sua digestão"],
      "checklist 'sua semana de digestão leve'",
      "visual fresco e saudável, tons de verde e amarelo, ícones de bem-estar, fácil de ler no celular",
      "a person feeling light after eating, digestive-friendly foods, herbal tea, a comfortable healthy belly",
      ["ALIMENTOS QUE PESAM", "MASTIGUE CERTO", "ADEUS AZIA", "DIGESTÃO LEVE"],
      "tones of green and yellow"),
]

# ===========================================================================
# MARKETING DIGITAL
# ===========================================================================
MARKETING = [
    E("Bio de Instagram que Vende", "Transforme Visitantes em Clientes em 1 Linha",
      "donos de negócio e criadores que têm Instagram mas perdem clientes por causa de uma bio fraca.",
      "por que sua bio decide se a pessoa vira cliente ou vai embora",
      ["O que sua bio precisa responder em 3 segundos", "A fórmula da bio que converte",
       "Nome e @ que aparecem na busca", "A chamada para ação que gera clique",
       "Link na bio: para onde mandar o cliente", "Destaques que vendem por você", "Os erros que afastam clientes"],
      "modelo 'monte sua bio que vende agora'",
      "visual moderno e jovem, tons de roxo, rosa e azul, ícones de Instagram, fácil de ler no celular",
      "an Instagram profile on a phone, a social media bio, a person checking their profile, engagement icons",
      ["3 SEGUNDOS DECISIVOS", "A FÓRMULA QUE CONVERTE", "CTA QUE GERA CLIQUE", "MODELO PRONTO"],
      "tones of purple, pink and blue"),
    E("WhatsApp que Vende", "O Guia Para Fechar Mais Vendas Pelo Zap",
      "lojistas, autônomos e vendedores que usam WhatsApp mas perdem vendas por não saber conduzir a conversa.",
      "por que o WhatsApp é sua maior ferramenta de vendas (e você desperdiça)",
      ["Configure um perfil profissional", "A primeira resposta que prende o cliente",
       "Conduza a conversa até a venda", "Quebre objeções sem pressão",
       "Catálogo, listas e respostas rápidas", "Follow-up que recupera vendas", "Os erros que espantam o cliente no zap"],
      "modelo 'scripts de WhatsApp para copiar'",
      "visual prático e moderno, tons de verde e azul, ícones de chat, fácil de ler no celular",
      "a WhatsApp business chat selling, a person closing a sale on phone, a product catalog, a happy customer",
      ["PERFIL PROFISSIONAL", "PRIMEIRA RESPOSTA", "QUEBRE OBJEÇÕES", "SCRIPTS PRONTOS"],
      "tones of green and blue"),
    E("Sua Primeira Campanha no Instagram Ads", "Do Medo ao Primeiro Cliente",
      "pequenos empreendedores que querem anunciar no Instagram mas têm medo de perder dinheiro.",
      "por que impulsionar do jeito errado é jogar dinheiro fora",
      ["Antes de anunciar: o que ter pronto", "Defina o público certo do seu anúncio",
       "Crie um anúncio que para o dedo", "Quanto investir para começar",
       "Acompanhe os números que importam", "Otimize sem gastar mais", "Os erros que queimam seu orçamento"],
      "checklist 'sua primeira campanha passo a passo'",
      "visual moderno e estratégico, tons de azul e roxo, ícones de anúncios, fácil de ler no celular",
      "a smartphone showing an Instagram ad, ad performance charts, a person setting up a campaign, target audience concept",
      ["O QUE TER PRONTO", "PÚBLICO CERTO", "ANÚNCIO QUE PARA O DEDO", "QUANTO INVESTIR"],
      "tones of blue and purple"),
    E("Canva para Quem Não é Designer", "Posts Profissionais em 10 Minutos",
      "empreendedores e criadores que querem criar artes bonitas sem saber design e sem pagar designer.",
      "por que você não precisa ser designer para ter posts lindos",
      ["Conheça o Canva em 5 minutos", "Use templates a seu favor",
       "Cores e fontes que combinam sempre", "Crie um padrão visual para sua marca",
       "Posts, stories e carrosséis profissionais", "Banco de imagens e elementos grátis", "Os erros de design que afastam seguidores"],
      "checklist 'crie seu primeiro post profissional hoje'",
      "visual criativo e colorido, tons de roxo e azul, ícones de design, fácil de ler no celular",
      "a person designing on Canva, colorful social media templates, a laptop with graphic design, creative posts",
      ["CANVA EM 5 MINUTOS", "TEMPLATES A SEU FAVOR", "CORES QUE COMBINAM", "PADRÃO DA MARCA"],
      "tones of purple and blue"),
    E("1000 Seguidores Reais", "Cresça no Instagram Sem Comprar Seguidor",
      "pessoas que querem crescer no Instagram de forma orgânica e atrair seguidores que realmente engajam.",
      "por que comprar seguidor destrói o seu perfil",
      ["Defina seu nicho e sua mensagem", "O perfil que faz a pessoa seguir",
       "Conteúdo que atrai seguidor certo", "Use Reels para alcançar mais gente",
       "Hashtags e legendas que ajudam", "Engaje e crie comunidade", "Os erros que travam seu crescimento"],
      "plano 'seus primeiros 1000 seguidores'",
      "visual jovem e vibrante, tons de rosa, roxo e azul, ícones de crescimento social, fácil de ler no celular",
      "a growing Instagram following, engagement notifications, a content creator, social media growth chart",
      ["DEFINA SEU NICHO", "PERFIL QUE FAZ SEGUIR", "REELS QUE ALCANÇAM", "CRIE COMUNIDADE"],
      "tones of pink, purple and blue"),
    E("Tráfego Pago para Iniciantes", "Atraia Clientes Todos os Dias com Anúncios",
      "empreendedores que querem usar anúncios pagos para ter um fluxo constante de clientes.",
      "a diferença entre quem espera o cliente e quem vai buscá-lo",
      ["O que é tráfego pago (sem complicar)", "Onde anunciar: Meta, Google e mais",
       "O caminho do clique até a venda", "Crie anúncios que convertem",
       "Orçamento e teste para iniciantes", "Métricas que mostram se está dando certo", "Os erros que fazem perder dinheiro"],
      "checklist 'comece a anunciar com segurança'",
      "visual estratégico e moderno, tons de azul e laranja, ícones de tráfego e funil, fácil de ler no celular",
      "digital ads on screens, a marketing funnel, traffic and clicks concept, a person managing campaigns",
      ["ONDE ANUNCIAR", "DO CLIQUE À VENDA", "ANÚNCIOS QUE CONVERTEM", "MÉTRICAS QUE IMPORTAM"],
      "tones of blue and orange"),
    E("Copywriting que Vende", "As Palavras que Fazem o Cliente Comprar",
      "empreendedores, criadores e vendedores que querem escrever textos que convencem e vendem.",
      "por que as palavras certas valem mais que o melhor produto",
      ["O que é copywriting e por que funciona", "Headlines que prendem a atenção",
       "Desperte desejo falando da dor e do sonho", "Gatilhos mentais usados com ética",
       "A chamada para ação irresistível", "Estrutura de um texto que vende", "Os erros de escrita que afastam a venda"],
      "modelos 'fórmulas de copy para copiar'",
      "visual sofisticado e moderno, tons de roxo e azul, ícones de escrita, fácil de ler no celular",
      "persuasive text on a screen, a copywriter at work, a sales page, words and conversion concept",
      ["HEADLINES QUE PRENDEM", "DOR E DESEJO", "GATILHOS MENTAIS", "CTA IRRESISTÍVEL"],
      "tones of purple and blue"),
    E("Funil de Vendas Simples", "Transforme Curiosos em Clientes Fiéis",
      "empreendedores digitais iniciantes que ouvem falar de 'funil' mas não sabem montar o seu.",
      "por que vender sem funil é desperdiçar clientes",
      ["O que é um funil de vendas na prática", "Atraia: como gerar visitantes",
       "Capture: transforme visita em contato", "Relacione: ganhe a confiança",
       "Venda: a oferta na hora certa", "Fidelize: faça o cliente voltar", "Os erros que furam seu funil"],
      "mapa 'desenhe seu funil simples'",
      "visual estratégico e clean, tons de azul e roxo, ícones de funil, fácil de ler no celular",
      "a sales funnel diagram, customer journey concept, a person planning strategy, conversion steps",
      ["ATRAIR VISITANTES", "CAPTURAR CONTATOS", "GANHAR CONFIANÇA", "VENDA NA HORA CERTA"],
      "tones of blue and purple"),
    E("E-mail Marketing do Zero", "Venda no Automático Para Sua Lista",
      "criadores e pequenos negócios que querem construir uma lista de e-mails e vender de forma recorrente.",
      "por que sua lista de e-mails é o ativo mais valioso do seu negócio",
      ["Por que e-mail ainda vende muito", "Como construir sua lista do zero",
       "A isca que faz a pessoa se inscrever", "Escreva e-mails que são abertos",
       "Sequências que vendem no automático", "Frequência e relacionamento certos", "Os erros que mandam você para o spam"],
      "modelo 'sua primeira sequência de e-mails'",
      "visual clean e profissional, tons de azul e roxo, ícones de e-mail, fácil de ler no celular",
      "an email inbox on a screen, a person writing emails, an email list growing, automation concept",
      ["CONSTRUA SUA LISTA", "A ISCA PERFEITA", "E-MAILS QUE ABREM", "VENDA NO AUTOMÁTICO"],
      "tones of blue and purple"),
    E("Marketing para Negócio Local", "Atraia Mais Clientes na Sua Cidade",
      "donos de comércios e serviços locais (salão, restaurante, loja) que querem atrair mais clientes da região.",
      "por que o cliente que está perto não está te encontrando",
      ["Apareça no Google da sua cidade", "Use o Google Meu Negócio a seu favor",
       "Instagram local que atrai a vizinhança", "Promoções que enchem a loja",
       "Peça e use avaliações de clientes", "Parcerias locais que multiplicam", "Os erros que afastam o cliente da região"],
      "checklist 'seu plano de marketing local'",
      "visual prático e acolhedor, tons de azul e laranja, ícones de loja e localização, fácil de ler no celular",
      "a local shop with customers, a Google Maps pin, a small business owner, a busy local store",
      ["APAREÇA NO GOOGLE", "GOOGLE MEU NEGÓCIO", "INSTAGRAM LOCAL", "PLANO LOCAL"],
      "tones of blue and orange"),
]

# ===========================================================================
# CULINÁRIA
# ===========================================================================
CULINARIA = [
    E("Marmitas Fit para a Semana", "Coma Bem, Economize e Emagreça",
      "pessoas que querem comer saudável durante a semana, economizar tempo e dinheiro com marmitas práticas.",
      "por que marmita fit é o segredo de quem come bem sem gastar muito",
      ["Planeje o cardápio da semana", "Lista de compras inteligente e barata",
       "Prepare tudo em 2 horas (batch cooking)", "Proteínas, carbo e legumes equilibrados",
       "Como armazenar e congelar certo", "Marmitas que não enjoam", "Os erros que estragam suas marmitas"],
      "plano 'monte suas marmitas da semana'",
      "visual fresco e organizado, tons de verde e laranja, fotos de comida saudável, fácil de ler no celular",
      "meal prep containers with healthy food, organized weekly marmitas, fresh ingredients, a tidy kitchen",
      ["CARDÁPIO DA SEMANA", "COMPRAS BARATAS", "PREPARE EM 2 HORAS", "CONGELE CERTO"],
      "tones of green and orange"),
    E("Sobremesas Sem Açúcar", "Doces Deliciosos Que Não Saem da Dieta",
      "pessoas que amam doce mas querem cortar o açúcar sem abrir mão da sobremesa.",
      "por que você pode comer doce e ainda cuidar da saúde",
      ["Os melhores substitutos do açúcar", "Brigadeiro fit que engana o paladar",
       "Mousses cremosos sem culpa", "Bolos e tortas sem açúcar refinado",
       "Sorvetes e geladinhos saudáveis", "Doces com frutas que adoçam naturalmente", "Os erros que estragam a sobremesa fit"],
      "lista '5 sobremesas sem açúcar para hoje'",
      "visual apetitoso e clean, tons de rosa, marrom e verde, fotos de doces saudáveis, fácil de ler no celular",
      "healthy sugar-free desserts, a fit brigadeiro, a creamy mousse, colorful fruit sweets",
      ["SUBSTITUTOS DO AÇÚCAR", "BRIGADEIRO FIT", "MOUSSES SEM CULPA", "5 RECEITAS PRONTAS"],
      "tones of pink, brown and green"),
    E("Receitas com 5 Ingredientes", "Pratos Deliciosos Sem Complicação",
      "pessoas sem tempo ou paciência que querem cozinhar pratos gostosos com poucos ingredientes.",
      "por que cozinhar pode ser simples, rápido e gostoso",
      ["O básico que não pode faltar na despensa", "Cafés da manhã com 5 ingredientes",
       "Almoços rápidos e completos", "Jantares práticos para o dia a dia",
       "Lanches que salvam a fome", "Sobremesas fáceis com poucos itens", "Os erros que complicam a cozinha"],
      "lista '10 receitas de 5 ingredientes'",
      "visual prático e apetitoso, tons de laranja e amarelo, fotos de comida simples, fácil de ler no celular",
      "simple dishes with few ingredients, a quick home-cooked meal, fresh basic ingredients, an easy recipe scene",
      ["DESPENSA BÁSICA", "ALMOÇOS RÁPIDOS", "JANTARES PRÁTICOS", "10 RECEITAS FÁCEIS"],
      "tones of orange and yellow"),
    E("Pães Caseiros para Iniciantes", "Faça Pão Fresquinho em Casa Sem Erro",
      "pessoas que querem aprender a fazer pão caseiro do zero, mesmo sem nenhuma experiência.",
      "por que pão caseiro é mais fácil (e gostoso) do que você imagina",
      ["Os ingredientes e utensílios básicos", "O segredo do fermento que dá certo",
       "Pão caseiro tradicional passo a passo", "Pão de forma macio para o dia a dia",
       "Pãezinhos e baguetes simples", "Como conservar o pão fresco por mais tempo", "Os erros que fazem o pão não crescer"],
      "lista '3 pães para você fazer esta semana'",
      "visual aconchegante e artesanal, tons de marrom e dourado, fotos de pães, fácil de ler no celular",
      "fresh homemade bread, a baker kneading dough, golden loaves on a table, a cozy kitchen with bread",
      ["O FERMENTO CERTO", "PÃO TRADICIONAL", "PÃO DE FORMA MACIO", "3 PÃES PRONTOS"],
      "tones of brown and gold"),
    E("Air Fryer", "30 Receitas Crocantes Sem Óleo",
      "donos de air fryer que querem aproveitar o aparelho ao máximo com receitas práticas e saudáveis.",
      "por que a air fryer virou a queridinha da cozinha",
      ["Conheça sua air fryer e os tempos certos", "Salgados crocantes sem fritura",
       "Carnes e frango suculentos", "Legumes e batatas perfeitos",
       "Lanches e petiscos rápidos", "Sobremesas na air fryer", "Os erros que ressecam ou queimam a comida"],
      "lista '10 receitas de air fryer para começar'",
      "visual moderno e apetitoso, tons de vermelho e laranja, fotos de comida crocante, fácil de ler no celular",
      "crispy air fryer food, golden fries and chicken, an air fryer on a counter, delicious healthy snacks",
      ["TEMPOS CERTOS", "CROCANTE SEM ÓLEO", "CARNES SUCULENTAS", "10 RECEITAS PRONTAS"],
      "tones of red and orange"),
    E("Receitas Low Carb", "Coma Bem e Emagreça Sem Passar Fome",
      "pessoas que querem reduzir carboidratos para emagrecer comendo comida de verdade e saborosa.",
      "como comer low carb sem sofrimento e sem enjoar",
      ["O que é low carb (de forma simples)", "Cafés da manhã low carb que saciam",
       "Almoços e jantares saborosos", "Substitua o pão, o arroz e a massa",
       "Lanches low carb para a fome", "Sobremesas low carb sem culpa", "Os erros que travam o emagrecimento low carb"],
      "lista '10 receitas low carb para a semana'",
      "visual fresco e saudável, tons de verde e marrom, fotos de comida low carb, fácil de ler no celular",
      "low carb healthy meals, fresh proteins and vegetables, a colorful low carb plate, a fit healthy dish",
      ["O QUE É LOW CARB", "SUBSTITUA PÃO E ARROZ", "LANCHES QUE SACIAM", "10 RECEITAS PRONTAS"],
      "tones of green and brown"),
    E("Doces para Vender e Lucrar", "Comece sua Renda Extra na Cozinha",
      "pessoas que querem ganhar dinheiro vendendo doces caseiros, do zero, mesmo sem experiência.",
      "por que doces caseiros são uma das rendas extras mais lucrativas",
      ["Os doces que mais vendem", "Calcule o preço e o lucro certo",
       "Brigadeiros gourmet que encantam", "Bolos de pote e no copo",
       "Embalagem que valoriza e vende mais", "Como divulgar e conseguir clientes", "Os erros que dão prejuízo"],
      "checklist 'comece a vender doces esta semana'",
      "visual apetitoso e comercial, tons de rosa e dourado, fotos de doces, fácil de ler no celular",
      "gourmet brigadeiros, beautifully packaged sweets, a person selling desserts, colorful cake pots",
      ["DOCES QUE MAIS VENDEM", "PREÇO E LUCRO", "BRIGADEIRO GOURMET", "COMECE A VENDER"],
      "tones of pink and gold"),
    E("Cozinha para Solteiros", "Receitas Práticas Para Quem Mora Sozinho",
      "pessoas que moram sozinhas e querem parar de pedir delivery, comendo bem e gastando pouco.",
      "por que cozinhar para um pode ser fácil, barato e gostoso",
      ["Monte uma despensa esperta para 1 pessoa", "Receitas rápidas para o dia a dia",
       "Cozinhe uma vez, coma a semana toda", "Pratos baratos e nutritivos",
       "Aproveite tudo e evite desperdício", "Refeições para impressionar uma visita", "Os erros de quem cozinha sozinho"],
      "lista '10 receitas práticas para 1 pessoa'",
      "visual descontraído e prático, tons de azul e laranja, fotos de comida simples, fácil de ler no celular",
      "a single-portion meal, a person cooking alone happily, simple practical dishes, a cozy small kitchen",
      ["DESPENSA ESPERTA", "COZINHE UMA VEZ", "PRATOS BARATOS", "10 RECEITAS PARA 1"],
      "tones of blue and orange"),
    E("Bolos Caseiros que Vendem", "Receitas Lucrativas Para Sua Confeitaria",
      "pessoas que querem fazer e vender bolos caseiros, transformando a cozinha em fonte de renda.",
      "por que um bom bolo caseiro nunca falta cliente",
      ["Os bolos campeões de venda", "A massa fofinha que dá certo sempre",
       "Recheios e coberturas irresistíveis", "Bolo no pote: o queridinho que vende muito",
       "Preço, lucro e produção em escala", "Apresentação e divulgação que vendem", "Os erros que estragam o bolo"],
      "checklist 'comece sua confeitaria caseira'",
      "visual apetitoso e caseiro, tons de marrom e rosa, fotos de bolos, fácil de ler no celular",
      "homemade cakes, a slice of fluffy cake, cake pots, a person decorating a cake",
      ["BOLOS QUE VENDEM", "MASSA FOFINHA", "BOLO NO POTE", "PREÇO E LUCRO"],
      "tones of brown and pink"),
    E("Receitas Veganas Práticas", "Comida Vegana Saborosa Para o Dia a Dia",
      "pessoas que querem comer mais vegano (ou começar) com receitas simples, baratas e gostosas.",
      "por que comida vegana não é cara nem complicada",
      ["O básico da cozinha vegana", "Proteínas vegetais que saciam",
       "Cafés da manhã veganos práticos", "Almoços e jantares saborosos",
       "Substitua ovo, leite e queijo", "Sobremesas veganas deliciosas", "Os erros de quem está começando no vegano"],
      "lista '10 receitas veganas para a semana'",
      "visual fresco e colorido, tons de verde e amarelo, fotos de comida vegana, fácil de ler no celular",
      "colorful vegan dishes, fresh plant-based ingredients, a vibrant vegan plate, a healthy green meal",
      ["O BÁSICO VEGANO", "PROTEÍNAS VEGETAIS", "SUBSTITUA LEITE E OVO", "10 RECEITAS PRONTAS"],
      "tones of green and yellow"),
]

TEMAS = {
    "Financas": ("Finanças", FINANCAS),
    "Saude": ("Saúde & Bem-estar", SAUDE),
    "Marketing": ("Marketing Digital", MARKETING),
    "Culinaria": ("Culinária", CULINARIA),
}

# ---------------------------------------------------------------------------
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def add_code_block(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def gerar_doc(nome_tema, ebooks, arquivo):
    doc = Document()
    doc.add_heading(f"Micro Guias — 10 Ebooks de {nome_tema}", level=0)
    p = doc.add_paragraph()
    r = p.add_run("Prompts prontos para o Gamma (conteúdo) + prompts de capa no padrão da coleção. "
                  "Para cada ebook: copie o bloco do Gamma e o bloco da capa.")
    r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
    doc.add_paragraph()

    for i, e in enumerate(ebooks, 1):
        h = doc.add_heading(f"{i}. {e['titulo']}", level=1)
        for run in h.runs:
            run.font.color.rgb = INDIGO
        sub = doc.add_paragraph()
        rr = sub.add_run(e["subtitulo"]); rr.bold = True; rr.font.color.rgb = GRAY
        doc.add_heading("▸ Prompt do ebook (Gamma)", level=2)
        add_code_block(doc, ebook_prompt(e))
        doc.add_heading("▸ Prompt da capa (IA de imagem)", level=2)
        add_code_block(doc, capa_prompt(e))
        if i < len(ebooks):
            doc.add_page_break()

    out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", arquivo))
    doc.save(out)
    print("Gerado:", out)


for chave, (nome, ebooks) in TEMAS.items():
    gerar_doc(nome, ebooks, f"Micro-Guias-10-Ebooks-{chave}.docx")
