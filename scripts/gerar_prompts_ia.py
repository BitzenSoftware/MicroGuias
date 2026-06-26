# -*- coding: utf-8 -*-
"""Gera um .docx com 10 prompts de ebook de IA (nichados) + prompts de capa."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Dados dos 10 ebooks
# Cada item: titulo, subtitulo, publico, intro, capitulos[7], conclusao,
#            design (cores do ebook), fotos (capa), teasers[4], cores (capa)
# ---------------------------------------------------------------------------
EBOOKS = [
    {
        "titulo": "ChatGPT para Pequenos Negócios",
        "subtitulo": "Automatize Tarefas e Atenda Melhor",
        "publico": "donos de pequenos negócios e autônomos brasileiros que querem usar IA para ganhar tempo, sem saber nada de tecnologia.",
        "intro": "como a IA pode virar seu funcionário mais barato e eficiente",
        "capitulos": [
            "Respostas prontas para clientes em segundos",
            "Crie descrições de produtos que vendem",
            "Planeje promoções e datas comemorativas",
            "Escreva e-mails e mensagens profissionais",
            "Gere ideias de conteúdo para divulgar",
            "Organize finanças e controle de estoque com ajuda da IA",
            "Os erros que fazem a IA te dar respostas inúteis",
        ],
        "conclusao": "lista '10 comandos prontos para o seu negócio hoje'",
        "design": "visual moderno e profissional, tons de azul e verde, ícones de negócios e tecnologia, espaço em branco, fácil de ler no celular",
        "fotos": "a small business owner using a laptop with AI, a phone showing automated customer messages, a shop counter with a tablet, a happy entrepreneur working",
        "teasers": ["ATENDIMENTO AUTOMÁTICO", "DESCRIÇÕES QUE VENDEM", "E-MAILS PROFISSIONAIS", "10 COMANDOS PRONTOS"],
        "cores_capa": "modern tech tones of blue and green, with refined golden accents",
    },
    {
        "titulo": "IA que Cria Imagens",
        "subtitulo": "Do Texto à Arte Profissional Sem Saber Desenhar",
        "publico": "pessoas criativas, empreendedores e criadores de conteúdo que querem gerar imagens incríveis com IA, mesmo sem nenhuma experiência em design.",
        "intro": "o que é IA de imagem e por que ela mudou tudo para quem cria",
        "capitulos": [
            "As melhores ferramentas gratuitas para começar",
            "A fórmula de um bom prompt de imagem",
            "Crie logos e identidade visual do zero",
            "Gere fotos de produtos e mockups",
            "Arte para redes sociais e thumbnails",
            "Estilos, iluminação e detalhes que fazem diferença",
            "Erros comuns que estragam suas imagens",
        ],
        "conclusao": "lista '10 prompts de imagem para copiar e testar'",
        "design": "visual artístico e moderno, tons de roxo, rosa e azul, elementos criativos e coloridos, fácil de ler no celular",
        "fotos": "a person creating digital art with AI on a screen, colorful AI-generated images, a creative workspace with a tablet, abstract digital artwork",
        "teasers": ["FERRAMENTAS GRÁTIS", "LOGOS DO ZERO", "FOTOS DE PRODUTOS", "10 PROMPTS PRONTOS"],
        "cores_capa": "vibrant tones of purple, pink and blue, with refined golden accents",
    },
    {
        "titulo": "Prompts que Funcionam",
        "subtitulo": "O Guia para Pedir Certo e Receber Respostas Incríveis",
        "publico": "qualquer pessoa que usa ChatGPT ou outras IAs mas se frustra com respostas ruins e quer dominar a arte de conversar com a IA.",
        "intro": "por que a IA te dá respostas ruins (e a culpa não é dela)",
        "capitulos": [
            "A anatomia de um prompt perfeito",
            "Dê contexto: o segredo que muda tudo",
            "Peça formato: listas, tabelas, passo a passo",
            "Use exemplos para guiar a resposta",
            "Refine e corrija sem começar do zero",
            "Personas: faça a IA agir como especialista",
            "Os 7 erros que arruínam seus prompts",
        ],
        "conclusao": "lista '15 modelos de prompt para copiar e adaptar'",
        "design": "visual limpo e tecnológico, tons de azul e roxo, ícones de chat e balões de conversa, muito espaço em branco",
        "fotos": "a chat conversation on a screen, a person typing on a laptop, glowing AI interface, organized text and lists on a display",
        "teasers": ["O PROMPT PERFEITO", "O PODER DO CONTEXTO", "PERSONAS DE IA", "15 MODELOS PRONTOS"],
        "cores_capa": "modern tones of blue and purple, with refined golden accents",
    },
    {
        "titulo": "IA para Estudantes",
        "subtitulo": "Aprenda Mais Rápido e Tire Notas Melhores",
        "publico": "estudantes do ensino médio, faculdade e concurseiros que querem usar IA para estudar melhor, organizar conteúdo e render mais.",
        "intro": "como a IA pode virar seu professor particular 24 horas",
        "capitulos": [
            "Transforme qualquer matéria em resumo fácil",
            "Crie questões e simulados para treinar",
            "Entenda assuntos difíceis com explicações simples",
            "Monte cronogramas de estudo realistas",
            "Revise mais rápido com mapas mentais",
            "Melhore redações e trabalhos acadêmicos",
            "Como usar IA sem cair em respostas erradas",
        ],
        "conclusao": "lista '10 comandos para estudar melhor hoje'",
        "design": "visual jovem e organizado, tons de azul, verde e amarelo, ícones de educação e livros, fácil de ler no celular",
        "fotos": "a student studying with a laptop, books and a tablet showing summaries, a young person taking notes, an organized study desk",
        "teasers": ["RESUMOS AUTOMÁTICOS", "SIMULADOS NA HORA", "CRONOGRAMA DE ESTUDO", "10 COMANDOS PRONTOS"],
        "cores_capa": "fresh tones of blue, green and yellow, with refined golden accents",
    },
    {
        "titulo": "Ganhe Dinheiro com IA",
        "subtitulo": "10 Formas de Monetizar com Inteligência Artificial",
        "publico": "pessoas que querem uma renda extra ou um novo negócio usando IA, partindo do zero e sem precisar investir muito.",
        "intro": "por que a IA abriu uma das maiores oportunidades da década",
        "capitulos": [
            "Crie e venda artes e designs com IA",
            "Ofereça serviços de escrita e copywriting",
            "Produza conteúdo para redes de terceiros",
            "Monte ebooks e produtos digitais (como este)",
            "Edite vídeos e crie cortes com IA",
            "Preste consultoria de prompts e automação",
            "Os erros que fazem iniciantes desistirem",
        ],
        "conclusao": "lista 'por onde começar nos próximos 7 dias'",
        "design": "visual moderno e aspiracional, tons de verde, dourado e azul, ícones de dinheiro e crescimento, fácil de ler no celular",
        "fotos": "a person earning money online with a laptop, financial growth charts, a freelancer working from home, a phone showing online payment",
        "teasers": ["VENDA ARTES COM IA", "SERVIÇOS DE ESCRITA", "PRODUTOS DIGITAIS", "COMECE EM 7 DIAS"],
        "cores_capa": "tones of green, gold and blue, with refined golden accents",
    },
    {
        "titulo": "IA para Redes Sociais",
        "subtitulo": "Crie Conteúdo Todo Dia Sem Travar",
        "publico": "criadores de conteúdo, empreendedores e social media que querem usar IA para nunca mais ficar sem ideia de post.",
        "intro": "por que você trava na hora de postar (e como a IA resolve)",
        "capitulos": [
            "Gere 30 ideias de post em 1 minuto",
            "Escreva legendas que prendem e vendem",
            "Roteiros de Reels e vídeos curtos",
            "Crie carrosséis que as pessoas salvam",
            "Planeje um calendário de conteúdo completo",
            "Responda comentários e DMs mais rápido",
            "Os erros que fazem seu conteúdo não engajar",
        ],
        "conclusao": "lista '10 comandos para nunca mais faltar conteúdo'",
        "design": "visual jovem e vibrante, tons de roxo, rosa e azul, ícones de redes sociais, fácil de ler no celular",
        "fotos": "a content creator with a phone, social media feed on a screen, a person recording a video, colorful social media icons",
        "teasers": ["30 IDEIAS EM 1 MINUTO", "LEGENDAS QUE VENDEM", "ROTEIROS DE REELS", "CALENDÁRIO DE POSTS"],
        "cores_capa": "vibrant tones of purple, pink and blue, with refined golden accents",
    },
    {
        "titulo": "Automatize seu Trabalho com IA",
        "subtitulo": "Planilhas, E-mails e Relatórios no Automático",
        "publico": "profissionais e funcionários de escritório que querem usar IA para fazer mais rápido as tarefas chatas do trabalho.",
        "intro": "quanto tempo você perde com tarefas que a IA faria por você",
        "capitulos": [
            "Crie e organize planilhas em segundos",
            "Escreva e responda e-mails profissionais",
            "Transforme dados em relatórios prontos",
            "Resuma reuniões, atas e documentos longos",
            "Monte apresentações e textos corporativos",
            "Crie modelos e processos reutilizáveis",
            "Os cuidados ao usar IA no trabalho",
        ],
        "conclusao": "lista '10 comandos para produzir mais no trabalho'",
        "design": "visual profissional e limpo, tons de azul e cinza, ícones de produtividade e escritório, fácil de ler no celular",
        "fotos": "a professional working with spreadsheets and AI, an organized office desk, a laptop showing reports, a person in a productive workflow",
        "teasers": ["PLANILHAS NA HORA", "E-MAILS PROFISSIONAIS", "RELATÓRIOS PRONTOS", "10 COMANDOS PRONTOS"],
        "cores_capa": "professional tones of blue and slate gray, with refined golden accents",
    },
    {
        "titulo": "IA para Currículo e Entrevista",
        "subtitulo": "Conquiste o Emprego dos Seus Sonhos",
        "publico": "pessoas em busca de emprego ou recolocação que querem usar IA para se destacar e passar em entrevistas.",
        "intro": "por que a maioria dos currículos é descartada em segundos",
        "capitulos": [
            "Monte um currículo que passa nos filtros",
            "Escreva uma carta de apresentação matadora",
            "Otimize seu LinkedIn para ser encontrado",
            "Treine respostas para perguntas difíceis",
            "Pesquise a empresa e impressione na entrevista",
            "Negocie salário com argumentos certos",
            "Os erros que reprovam bons candidatos",
        ],
        "conclusao": "lista '10 comandos para conseguir a vaga'",
        "design": "visual profissional e confiante, tons de azul-marinho e dourado, ícones de carreira, fácil de ler no celular",
        "fotos": "a person in a job interview, a professional resume on a screen, a confident candidate shaking hands, a LinkedIn profile on a phone",
        "teasers": ["CURRÍCULO QUE PASSA", "LINKEDIN OTIMIZADO", "RESPOSTAS DIFÍCEIS", "NEGOCIE SALÁRIO"],
        "cores_capa": "professional tones of navy blue and gold, with refined golden accents",
    },
    {
        "titulo": "Organize sua Vida com IA",
        "subtitulo": "Assistentes Inteligentes para a Sua Rotina",
        "publico": "pessoas atarefadas que querem usar IA para organizar tarefas, finanças, casa e tempo de forma simples.",
        "intro": "como ter um assistente pessoal de graça no seu bolso",
        "capitulos": [
            "Planeje sua semana e suas metas",
            "Crie listas de tarefas e lembretes inteligentes",
            "Organize as finanças pessoais do mês",
            "Monte cardápios e listas de compras",
            "Planeje viagens e passeios do começo ao fim",
            "Crie hábitos e rotinas que funcionam",
            "Os erros que tiram o foco da IA",
        ],
        "conclusao": "lista '10 comandos para organizar sua vida hoje'",
        "design": "visual leve e organizado, tons de azul, verde-água e branco, ícones de organização e calendário, fácil de ler no celular",
        "fotos": "a person organizing tasks on a phone, a clean planner and calendar, an organized home desk, a relaxed person managing routine",
        "teasers": ["PLANEJE A SEMANA", "FINANÇAS DO MÊS", "CARDÁPIO E COMPRAS", "10 COMANDOS PRONTOS"],
        "cores_capa": "light tones of blue, teal and white, with refined golden accents",
    },
    {
        "titulo": "IA para Quem Vende",
        "subtitulo": "Atendimento, Copy e Fechamento com Inteligência Artificial",
        "publico": "vendedores, lojistas e empreendedores que querem usar IA para vender mais, atender melhor e fechar mais negócios.",
        "intro": "como a IA pode dobrar suas vendas sem aumentar sua equipe",
        "capitulos": [
            "Crie scripts de venda que convertem",
            "Responda objeções com argumentos certos",
            "Escreva anúncios e copy que vendem",
            "Personalize o atendimento de cada cliente",
            "Faça follow-up sem ser chato",
            "Analise e melhore suas vendas com IA",
            "Os erros que espantam o cliente",
        ],
        "conclusao": "lista '10 comandos para vender mais ainda hoje'",
        "design": "visual energético e profissional, tons de laranja, vermelho e azul, ícones de vendas e crescimento, fácil de ler no celular",
        "fotos": "a salesperson closing a deal, a phone showing a sales conversation, a confident seller with a customer, rising sales charts",
        "teasers": ["SCRIPTS QUE CONVERTEM", "QUEBRE OBJEÇÕES", "COPY QUE VENDE", "10 COMANDOS PRONTOS"],
        "cores_capa": "energetic tones of orange, red and blue, with refined golden accents",
    },
]

# ---------------------------------------------------------------------------
# Templates
# ---------------------------------------------------------------------------
def ebook_prompt(e):
    caps = "\n".join(f"{i+3}. Capítulo {i+1} — {c}" for i, c in enumerate(e["capitulos"]))
    return f'''Crie um ebook prático e direto chamado "{e['titulo']}: {e['subtitulo']}".

Público-alvo: {e['publico']}

Tom: motivador, simples e direto. Linguagem fácil (português do Brasil), frases curtas, zero termos técnicos. Trate o leitor por "você".

Formato: ebook de 12 a 15 páginas, objetivo e acionável. Cada capítulo cabe em 1 página, com no máximo 3 parágrafos curtos, e termina com um quadro destacado "Comando pronto" — um exemplo de prompt ou ação que o leitor pode copiar e aplicar na hora.

Estrutura:
1. Capa atraente com o título e subtítulo
2. Introdução: {e['intro']}
{caps}
10. Conclusão + {e['conclusao']}

Design: {e['design']}. Inclua um quadro de destaque "Comando pronto" em cada capítulo.

Adicione no rodapé: "Micro Guias".'''


def capa_prompt(e):
    t = e["teasers"]
    return f'''Create ONE single professional 3D ebook cover mockup. Output must contain ONLY ONE book — never multiple books, never variations, never a collage.

BOOK / FORMAT:
- A single hardcover book of normal thickness, standing upright in 3/4 perspective, rotated about 20 degrees to the RIGHT (spine and page edges visible on the RIGHT side)
- The ENTIRE book and the FULL title fully visible and centered, with comfortable empty margin — do NOT crop anything
- Clean solid white background, soft realistic shadow and a subtle reflection below
- Photorealistic, high quality, premium luxury collection style

COVER DESIGN (top to bottom):
1. A top pill badge with a thin golden outline containing the text "EDIÇÃO ESPECIAL" in light letters
2. A clean, organized 2x2 GRID of four separate rectangular photos, each with a thin golden border, aligned in perfect rows and columns — NOT floating, NOT overlapping, NOT scattered. The four photos: {e['fotos']}
3. A central rounded banner in DARK deep navy blue color, framed by a thin elegant GOLDEN ornamental border; inside it the main title in a refined golden serif font
4. Just below, the subtitle in a clean light/golden font
5. A dark bottom band with 4 short teasers in small golden uppercase letters, separated by dots

TEXTS (exactly, in Portuguese):
- Title: "{e['titulo']}"
- Subtitle: "{e['subtitulo']}"
- Teasers: "{t[0]}" • "{t[1]}" • "{t[2]}" • "{t[3]}"

COLORS: {e['cores_capa']}; the title banner and bottom band are always dark navy with golden details. Theme accent colors only inside the photos.

OUTPUT: a single vertical book cover, 2:3 aspect ratio, one book only, centered, full book visible, with a CLEAN ALIGNED 2x2 photo grid, premium and elegant.'''


# ---------------------------------------------------------------------------
# Monta o documento
# ---------------------------------------------------------------------------
doc = Document()

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
GRAY = RGBColor(0x6B, 0x72, 0x80)

titulo = doc.add_heading("Micro Guias — 10 Ebooks de Inteligência Artificial", level=0)

p = doc.add_paragraph()
run = p.add_run("Prompts prontos para gerar no Gamma + prompts de capa (padrão da coleção). "
                "Para cada ebook, copie o bloco do Gamma para gerar o conteúdo e o bloco da capa "
                "para gerar a imagem numa IA de imagem (Gemini, Ideogram, etc.).")
run.italic = True
run.font.color.rgb = GRAY
run.font.size = Pt(10)

doc.add_paragraph()

def add_code_block(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    # leve recuo
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)

for i, e in enumerate(EBOOKS, 1):
    h = doc.add_heading(f"{i}. {e['titulo']}", level=1)
    for run in h.runs:
        run.font.color.rgb = INDIGO

    sub = doc.add_paragraph()
    r = sub.add_run(e["subtitulo"])
    r.bold = True
    r.font.color.rgb = GRAY

    doc.add_heading("▸ Prompt do ebook (Gamma)", level=2)
    add_code_block(ebook_prompt(e))

    doc.add_heading("▸ Prompt da capa (IA de imagem)", level=2)
    add_code_block(capa_prompt(e))

    if i < len(EBOOKS):
        doc.add_page_break()

import os
out_dir = os.path.join(os.path.dirname(__file__), "..", "..")
out_path = os.path.abspath(os.path.join(out_dir, "Micro-Guias-10-Ebooks-IA.docx"))
doc.save(out_path)
print("Arquivo gerado em:", out_path)
