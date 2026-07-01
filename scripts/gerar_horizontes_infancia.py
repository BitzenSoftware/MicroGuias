# -*- coding: utf-8 -*-
"""Gera o .docx da coleção 'Horizontes da Infância': 10 ebooks (20 páginas)
com prompt do Gamma + prompt de capa no padrão da loja."""
import os
from docx import Document
from docx.shared import Pt, RGBColor

COLECAO = "Horizontes da Infância"

# ---------------------------------------------------------------------------
def ebook_prompt(e):
    caps = e["capitulos"]
    linhas = "\n".join(f"{i+3}. Capítulo {i+1} — {c}" for i, c in enumerate(caps))
    n_conc = len(caps) + 3
    return f'''Crie um ebook profissional e prático chamado "{e['titulo']}: {e['subtitulo']}", da coleção "{COLECAO}".

Público-alvo: {e['publico']}

Tom: profissional, acolhedor e prático. Português do Brasil, linguagem clara e acessível, com embasamento mas sem academicismo pesado. Trate o leitor por "você". Conecte a teoria com a prática do dia a dia da Educação Infantil, com exemplos concretos de sala.

Formato: ebook de 20 páginas, bem estruturado e acionável. Cada capítulo ocupa cerca de 2 páginas, com 2 a 4 parágrafos objetivos, e termina com um quadro destacado "{e['destaque']}" — uma aplicação prática imediata para a sala ou a rotina.

Estrutura:
1. Capa com o título e o subtítulo
2. Introdução: {e['intro']}
{linhas}
{n_conc}. Conclusão + {e['conclusao']}

Design: {e['design']}. Inclua um quadro "{e['destaque']}" em cada capítulo e, quando fizer sentido, conexões leves com a BNCC (Campos de Experiência da Educação Infantil).

Adicione no rodapé: "Micro Guias — Coleção {COLECAO}".'''


def capa_prompt(e):
    t = e["teasers"]
    return f'''Create ONE single professional 3D ebook cover mockup. Output must contain ONLY ONE book — never multiple books, never variations, never a collage.

BOOK / FORMAT:
- A single hardcover book of normal thickness, standing upright in 3/4 perspective, rotated about 20 degrees to the RIGHT (spine and page edges visible on the RIGHT side)
- The ENTIRE book and the FULL title fully visible and centered, with comfortable empty margin — do NOT crop anything
- Clean solid white background, soft realistic shadow and a subtle reflection below
- Photorealistic, high quality, premium luxury collection style

COVER DESIGN (top to bottom):
1. A top pill badge with a thin golden outline containing the text "HORIZONTES DA INFÂNCIA" in light letters
2. A clean, organized 2x2 GRID of four separate rectangular photos, each with a thin golden border, aligned in perfect rows and columns — NOT floating, NOT overlapping, NOT scattered. The four photos: {e['fotos']}
3. A central rounded banner in DARK deep navy blue color, framed by a thin elegant GOLDEN ornamental border; inside it the main title in a refined golden serif font
4. Just below, the subtitle in a clean light/golden font
5. A dark bottom band with 4 short teasers in small golden uppercase letters, separated by dots

TEXTS (exactly, in Portuguese):
- Title: "{e['titulo']}"
- Subtitle: "{e['subtitulo']}"
- Teasers: "{t[0]}" • "{t[1]}" • "{t[2]}" • "{t[3]}"

COLORS: {e['cores_capa']}; the title banner and bottom band are always dark navy with golden details. Theme accent colors only inside the photos.

OUTPUT: a single vertical book cover, 2:3 aspect ratio, one book only, centered, full book visible, with a CLEAN ALIGNED 2x2 photo grid, premium and elegant.'''


def E(titulo, subtitulo, publico, intro, capitulos, conclusao, design, fotos, teasers, cores_capa, destaque):
    return dict(titulo=titulo, subtitulo=subtitulo, publico=publico, intro=intro,
                capitulos=capitulos, conclusao=conclusao, design=design, fotos=fotos,
                teasers=teasers, cores_capa=cores_capa, destaque=destaque)


EDUCADORES = ("professoras e professores da Educação Infantil, coordenadores pedagógicos "
              "e estudantes de Pedagogia")

# ===========================================================================
EBOOKS = [
    E("O Cérebro que Aprende", "Neurociência e Desenvolvimento na Primeira Infância",
      f"{EDUCADORES} que querem entender como a criança de 0 a 5 anos realmente aprende.",
      "como o cérebro infantil se desenvolve e por que os primeiros anos são decisivos para a vida toda",
      ["Como o cérebro da criança se forma (as janelas de oportunidade)",
       "Marcos do desenvolvimento cognitivo dos 0 aos 5 anos",
       "Desenvolvimento motor: do corpo à aprendizagem",
       "Desenvolvimento emocional e o vínculo do apego",
       "O papel das experiências e do ambiente no cérebro",
       "Como o afeto potencializa a aprendizagem",
       "Sinais de alerta no desenvolvimento: quando observar",
       "Práticas simples que estimulam o cérebro que aprende"],
      "um resumo visual dos marcos de desenvolvimento por faixa etária",
      "visual profissional e acolhedor, tons de azul e verde, ícones de cérebro e infância, ilustrações leves, fácil de ler no celular",
      "a curious young child exploring and learning, a soft colorful brain development illustration, a teacher gently observing a toddler, children playing and discovering together",
      ["MARCOS DOS 0 AOS 5 ANOS", "JANELAS DE APRENDIZAGEM", "AFETO E COGNIÇÃO", "SINAIS DE ALERTA"],
      "tones of blue and green", "Na prática"),

    E("A Pedagogia do Brincar", "O Lúdico como Eixo da Aprendizagem",
      f"{EDUCADORES} que querem transformar o brincar em aprendizagem intencional e planejada.",
      "por que o brincar é a linguagem e o 'trabalho' da criança — e a base de toda aprendizagem",
      ["Brincar é coisa séria: o que a ciência mostra",
       "Brincadeira livre x brincadeira dirigida: quando usar cada uma",
       "Os tipos de brincar (simbólico, motor, de construção, de regras)",
       "Como planejar brincadeiras com intenção pedagógica",
       "O papel do professor: mediar sem interromper",
       "Brinquedos e materiais não estruturados",
       "Como observar e registrar o que a criança aprende brincando",
       "Brincadeiras por Campo de Experiência (BNCC)"],
      "um banco de brincadeiras organizado por objetivo de aprendizagem",
      "visual lúdico e colorido, tons de laranja e azul, ícones de brinquedos e movimento, fácil de ler no celular",
      "children playing freely and joyfully, wooden toys and building blocks, a teacher playing alongside kids, a bright outdoor play scene",
      ["LIVRE x DIRIGIDO", "BRINCAR COM INTENÇÃO", "MATERIAIS NÃO ESTRUTURADOS", "BANCO DE BRINCADEIRAS"],
      "tones of orange and blue", "Leve para a sala"),

    E("Primeiras Palavras e Letramento", "Linguagem, Oralidade e o Caminho para a Escrita",
      f"{EDUCADORES} que querem desenvolver a linguagem e introduzir o mundo escrito sem antecipar a alfabetização.",
      "como a linguagem floresce na infância e por que 'ler o mundo' vem antes de ler palavras",
      ["Como a linguagem se desenvolve dos 0 aos 5 anos",
       "A oralidade como base de toda a aprendizagem",
       "O poder da contação de histórias",
       "Como criar um ambiente alfabetizador sem pressão",
       "Consciência fonológica de forma lúdica",
       "Rabiscos, garatujas e a escrita espontânea",
       "Literatura infantil: como escolher e mediar",
       "Letramento e a BNCC na Educação Infantil"],
      "um roteiro para montar um cantinho de leitura irresistível",
      "visual afetivo e literário, tons de vermelho e amarelo, ícones de livros e balões de fala, fácil de ler no celular",
      "a teacher reading a storybook to attentive children, a cozy reading corner with cushions, a small child scribbling on paper, colorful stacked picture books",
      ["ORALIDADE PRIMEIRO", "CONTAÇÃO DE HISTÓRIAS", "AMBIENTE ALFABETIZADOR", "SEM ANTECIPAR"],
      "tones of red and warm yellow", "Na prática"),

    E("Emoções em Construção", "Educação Socioemocional na Primeira Infância",
      f"{EDUCADORES} que querem ajudar as crianças a reconhecer emoções, lidar com frustrações e desenvolver empatia.",
      "por que ensinar a sentir e a nomear emoções é tão importante quanto ler e contar",
      ["O que é educação socioemocional (e por que agora)",
       "Como as emoções funcionam no cérebro infantil",
       "Ajude a criança a nomear o que sente",
       "Birras e frustrações: o que fazer e o que evitar",
       "Autorregulação: ensinando a criança a se acalmar",
       "Empatia e habilidades sociais no dia a dia",
       "O educador como modelo emocional",
       "Atividades socioemocionais para a rotina"],
      "um kit de estratégias para os momentos difíceis do dia",
      "visual acolhedor e sereno, tons de lilás e azul, ícones de emoções e coração, fácil de ler no celular",
      "a child expressing a strong emotion, a teacher gently comforting a child, a colorful chart of emotion faces, children sharing and showing empathy",
      ["NOMEAR AS EMOÇÕES", "LIDAR COM BIRRAS", "AUTORREGULAÇÃO", "EMPATIA NA PRÁTICA"],
      "tones of lavender and blue", "Na prática"),

    E("Espaços, Tempos e Rotinas", "Ambiente e Rotina que Acolhem e Educam",
      f"{EDUCADORES} que querem organizar espaços e rotinas que promovem autonomia, segurança e bem-estar.",
      "por que o espaço é o 'terceiro educador' e a rotina é fonte de segurança para a criança",
      ["O ambiente como terceiro educador",
       "Organização da sala por cantos e áreas de interesse",
       "O pátio e os espaços externos como lugar de aprender",
       "A importância de uma rotina previsível",
       "Transições tranquilas entre uma atividade e outra",
       "Momentos de cuidado: sono, alimentação e higiene",
       "Autonomia: um ambiente que a criança usa sozinha",
       "Como planejar uma rotina equilibrada"],
      "um modelo de rotina diária por faixa etária",
      "visual clean e organizado, tons de verde e azul, ícones de espaço, relógio e organização, fácil de ler no celular",
      "an organized classroom with learning corners, children playing in an outdoor area, a calm cozy nap space, a colorful daily routine chart",
      ["AMBIENTE QUE EDUCA", "CANTOS E ÁREAS", "TRANSIÇÕES TRANQUILAS", "ROTINA EQUILIBRADA"],
      "tones of green and blue", "Organize já"),

    E("Inclusão e Diversidade na Prática", "Estratégias para Acolher Todas as Crianças",
      f"{EDUCADORES} que querem uma prática inclusiva para crianças com TEA, TDAH, deficiências e diferentes necessidades.",
      "o que é uma educação inclusiva de verdade — para além do discurso, no dia a dia da sala",
      ["Inclusão: princípios e o que diz a lei",
       "Conhecer para acolher: cada criança é única",
       "Estratégias para crianças com TEA",
       "Apoiando crianças com TDAH",
       "Deficiências: adaptações simples no dia a dia",
       "Desenho Universal para a Aprendizagem (DUA)",
       "Diversidade cultural, racial e de famílias",
       "Parceria com a família e a equipe multidisciplinar"],
      "um checklist para avaliar o quanto sua sala é inclusiva",
      "visual inclusivo e colorido, tons de azul e verde, ícones de diversidade e acolhimento, fácil de ler no celular",
      "a diverse group of children playing together, a teacher supporting a child with special needs, inclusive tactile classroom materials, children of different backgrounds smiling",
      ["TEA E TDAH NA SALA", "ADAPTAÇÕES DIÁRIAS", "DESENHO UNIVERSAL", "SALA INCLUSIVA"],
      "tones of blue and green", "Adapte na prática"),

    E("Arte, Som e Movimento", "Psicomotricidade, Música e Artes na Infância",
      f"{EDUCADORES} que querem explorar a expressão corporal, musical e artística com intencionalidade.",
      "por que arte, som e movimento são aprendizagem de verdade — e não apenas passatempo",
      ["O corpo que aprende: psicomotricidade na prática",
       "Movimento e coordenação no dia a dia",
       "Musicalização na primeira infância",
       "Cantigas, ritmos e os sons do cotidiano",
       "Artes visuais: o processo importa mais que o produto",
       "Materiais e técnicas para explorar sem medo",
       "Expressão, criatividade e imaginação",
       "Projetos que integram arte, som e movimento"],
      "um banco de atividades de expressão para a semana",
      "visual artístico e vibrante, tons de laranja, rosa e azul, ícones de arte, música e movimento, fácil de ler no celular",
      "children painting freely with their hands, kids playing simple musical instruments, a joyful movement and dance activity, colorful art materials on a table",
      ["PSICOMOTRICIDADE", "MUSICALIZAÇÃO", "O PROCESSO IMPORTA", "ARTE INTEGRADA"],
      "tones of orange, pink and blue", "Leve para a sala"),

    E("Documentação Pedagógica e Avaliação", "Avaliar Sem Notas: Portfólios e Registros",
      f"{EDUCADORES} que querem avaliar e documentar a aprendizagem na Educação Infantil com qualidade.",
      "como avaliar onde não existem notas — e por que documentar é dar visibilidade à criança",
      ["Avaliar na Educação Infantil: o que realmente muda",
       "A observação: o instrumento mais poderoso do educador",
       "Registros do dia a dia: o que e como anotar",
       "Portfólios que contam a história da criança",
       "Relatórios e pareceres descritivos com sensibilidade",
       "Documentação pedagógica que comunica",
       "Fotos, vídeos e evidências com ética e cuidado",
       "Da avaliação ao replanejamento"],
      "modelos prontos de parecer descritivo e de portfólio",
      "visual profissional e organizado, tons de azul e verde-água, ícones de documento e observação, fácil de ler no celular",
      "a teacher writing observation notes, a child's colorful portfolio, pedagogical documentation displayed on a wall, a teacher thoughtfully reviewing records",
      ["AVALIAR SEM NOTAS", "OBSERVAÇÃO E REGISTRO", "PORTFÓLIOS", "PARECERES DESCRITIVOS"],
      "tones of blue and teal", "Modelo pronto"),

    E("A Aliança Invisível: Escola e Família", "Parceria e Comunicação com as Famílias",
      f"{EDUCADORES} que querem construir uma relação de confiança e parceria com as famílias.",
      "por que uma parceria verdadeira entre escola e família muda o destino da criança",
      ["Por que escola e família precisam caminhar juntas",
       "A acolhida e a adaptação feitas com a família",
       "Comunicação diária que constrói confiança",
       "Reuniões pedagógicas que valem a pena",
       "Como dar retornos difíceis com cuidado e empatia",
       "Envolver as famílias na aprendizagem dos filhos",
       "Famílias diversas: acolher sem julgar",
       "Resolvendo conflitos e alinhando expectativas"],
      "um roteiro pronto de reunião e de comunicação com os pais",
      "visual acolhedor e humano, tons de verde e azul, ícones de família e escola de mãos dadas, fácil de ler no celular",
      "a teacher talking warmly with parents, a parent-teacher meeting, a parent dropping a child at school, a trusting conversation between adults",
      ["COMUNICAÇÃO DIÁRIA", "REUNIÕES EFICAZES", "RETORNOS DIFÍCEIS", "PARCERIA DE CONFIANÇA"],
      "tones of green and blue", "Na prática"),

    E("A Infância do Futuro", "Natureza, Telas e Metodologias Ativas",
      f"{EDUCADORES} que querem atualizar a prática com as tendências mais relevantes da Educação Infantil.",
      "para onde caminha a infância — e como preparar sua prática para o mundo que vem aí",
      ["O desemparedamento: a infância que volta para o ar livre",
       "O contato com a natureza como aprendizagem",
       "Telas e tecnologia: uso consciente e equilibrado",
       "O que a ciência diz sobre o tempo de tela",
       "Metodologias ativas na Educação Infantil",
       "Abordagens que inspiram (Reggio Emilia, Montessori, Pikler)",
       "Sustentabilidade e cidadania desde cedo",
       "Preparando a criança para um mundo em mudança"],
      "um manifesto prático pela infância do futuro",
      "visual moderno e natural, tons de verde e azul, ícones de natureza e tecnologia equilibrada, fácil de ler no celular",
      "children exploring nature outdoors, kids learning in a garden, a teacher guiding balanced use of a tablet, a bright active modern classroom",
      ["DESEMPAREDAMENTO", "TELAS CONSCIENTES", "METODOLOGIAS ATIVAS", "NATUREZA QUE ENSINA"],
      "tones of green and blue", "Repense a prática"),
]

# ---------------------------------------------------------------------------
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def add_code_block(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def gerar_doc(arquivo):
    doc = Document()
    doc.add_heading(f"Micro Guias — Coleção {COLECAO}", level=0)
    p = doc.add_paragraph()
    r = p.add_run("10 ebooks de 20 páginas para educadores da Educação Infantil. "
                  "Para cada edição: copie o bloco do Gamma (conteúdo) e o bloco da capa "
                  "(IA de imagem). As capas seguem o padrão da coleção — só mudam título, "
                  "fotos e cores.")
    r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
    doc.add_paragraph()

    for i, e in enumerate(EBOOKS, 1):
        h = doc.add_heading(f"Edição {i}: {e['titulo']}", level=1)
        for run in h.runs:
            run.font.color.rgb = INDIGO
        sub = doc.add_paragraph()
        rr = sub.add_run(e["subtitulo"]); rr.bold = True; rr.font.color.rgb = GRAY
        doc.add_heading("▸ Prompt do ebook (Gamma) — 20 páginas", level=2)
        add_code_block(doc, ebook_prompt(e))
        doc.add_heading("▸ Prompt da capa (IA de imagem)", level=2)
        add_code_block(doc, capa_prompt(e))
        if i < len(EBOOKS):
            doc.add_page_break()

    out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", arquivo))
    doc.save(out)
    print("Gerado:", out)


gerar_doc("Micro-Guias-Horizontes-da-Infancia.docx")
